"""
exports/views.py

Handles generation and export of IET membership application forms
(single and bulk) as DOCX/PDF/ZIP files.

DOCX generation uses docxtpl. PDF conversion uses LibreOffice headless
mode via subprocess, making this module fully compatible with Linux
environments (PythonAnywhere, Ubuntu VPS, etc.) with no dependency on
Windows-only libraries such as win32com or pythoncom.
"""

import io
import os

import uuid
import zipfile
from datetime import datetime

from django.conf import settings
from django.contrib.auth.decorators import user_passes_test
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, render

from docxtpl import DocxTemplate


from applications.models import Application, ApplicationStatus

from payments.models import Payment


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# Directory used for temporary DOCX/PDF files during export.
# Kept inside BASE_DIR/tmp_exports so it's easy to purge and doesn't
# clutter the project root.
EXPORT_TMP_DIR = os.path.join(settings.BASE_DIR, "tmp_exports")

# Path to the LibreOffice binary. On most Linux servers (including
# PythonAnywhere and Ubuntu) "soffice" is the correct executable name;
# override via Django settings if your environment differs.
LIBREOFFICE_BIN = getattr(settings, "LIBREOFFICE_BIN", "soffice")


def _ensure_tmp_dir():
    """Ensure the temporary export directory exists."""
    os.makedirs(EXPORT_TMP_DIR, exist_ok=True)


def _safe_remove(*paths):
    """Remove files if they exist, silently ignoring errors."""
    for path in paths:
        try:
            if path and os.path.exists(path):
                os.remove(path)
        except OSError:
            pass


def _convert_docx_to_pdf(docx_path, outdir):
    """
    Convert a DOCX file to PDF using LibreOffice headless mode.

    Returns the path to the generated PDF file.
    Raises RuntimeError with a descriptive message if conversion fails.
    """
    try: 
        result = subprocess.run(
            [
                LIBREOFFICE_BIN,
                "--headless",
                "--convert-to",
                "pdf",
                docx_path,
                "--outdir",
                outdir,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
    except FileNotFoundError:
        raise RuntimeError(
            f"LibreOffice executable '{LIBREOFFICE_BIN}' was not found on this "
            "server. Please install LibreOffice or configure LIBREOFFICE_BIN "
            "in Django settings."
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice PDF conversion timed out.")

    if result.returncode != 0:
        raise RuntimeError(
            "LibreOffice PDF conversion failed "
            f"(exit code {result.returncode}): {result.stderr.strip() or result.stdout.strip()}"
        )

    expected_pdf = os.path.join(
        outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
    )

    if not os.path.exists(expected_pdf):
        raise RuntimeError(
            "LibreOffice reported success but the expected PDF file was not "
            f"found at {expected_pdf}."
        )

    return expected_pdf


def _build_application_context(application):
    """Build the docxtpl context dict for a single application."""
    return {
        "first_name": application.first_name,
        "middle_name": application.middle_name or "",
        "last_name": application.last_name,
        "full_name": application.full_name,
        "nationality": application.nationality,
        "date_of_birth": application.date_of_birth.strftime("%d/%m/%Y"),
        "age": application.age,
        "email": application.email,
        "phone_number": application.phone_number,
        "address": application.address,
        "city": application.city,
        "university": application.university,
        "department": application.department,
        "course": application.course,
        "year_of_study": f"{application.year_of_study} year",
        "date_of_admission": application.admission_date_formatted,
        "year_of_graduation": application.year_of_graduation,
        "application_date": application.created_at.strftime("%d/%m/%Y"),
        "status": application.get_status_display(),
        "reference_number": (
            application.payment.reference_number
            if hasattr(application, "payment")
            else "N/A"
        ),
    }


# ---------------------------------------------------------------------------
# Single application export
# ---------------------------------------------------------------------------
@user_passes_test(lambda u: u.is_staff)
def export_application_form(request, pk):
    """
    Export application form as DOCX using a docxtpl template.
    No PDF conversion required.
    """

    application = get_object_or_404(Application, pk=pk)

    # Path to DOCX template
    template_path = os.path.join(
        settings.BASE_DIR,
        "templates",
        "exports",
        "application_template.docx"
    )

    # Check template exists
    if not os.path.exists(template_path):
        return HttpResponse(
            "Template file not found. Please create application_template.docx "
            "inside templates/exports/",
            status=500
        )

    # Create temporary directory
    export_dir = os.path.join(settings.BASE_DIR, "tmp_exports")
    os.makedirs(export_dir, exist_ok=True)

    file_id = uuid.uuid4()

    temp_docx = os.path.join(
        export_dir,
        f"application_{file_id}.docx"
    )

    try:
        # Load template
        template = DocxTemplate(template_path)

        # Prepare data
        context = {
            "first_name": application.first_name,
            "middle_name": application.middle_name or "",
            "last_name": application.last_name,

            "full_name": application.full_name,

            "nationality": application.nationality,

            "date_of_birth": (
                application.date_of_birth.strftime("%d/%m/%Y")
                if application.date_of_birth
                else ""
            ),

            "age": application.age,

            "email": application.email,

            "phone_number": application.phone_number,

            "address": application.address,

            "city": application.city,

            "university": application.university,

            "department": application.department,

            "course": application.course,

            "year_of_study": (
                f"{application.year_of_study} year"
                if application.year_of_study
                else ""
            ),

            "date_of_admission": application.admission_date_formatted,

            "year_of_graduation": application.year_of_graduation,

            "application_date": (
                application.created_at.strftime("%d/%m/%Y")
            ),

            "status": application.get_status_display(),

            "reference_number": (
                application.payment.reference_number
                if hasattr(application, "payment")
                else "N/A"
            ),
        }

        # Fill template placeholders
        template.render(context)

        # Save completed document
        template.save(temp_docx)


        # Read generated DOCX
        with open(temp_docx, "rb") as file:
            docx_data = file.read()


        # Return DOCX response
        response = HttpResponse(
            docx_data,
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            )
        )

        filename = (
            f"IET_Application_{application.full_name.replace(' ', '_')}.docx"
        )

        response["Content-Disposition"] = (
            f'attachment; filename="{filename}"'
        )

        return response


    except Exception as e:
        import traceback

        error_details = traceback.format_exc()

        return HttpResponse(
            f"""
            Error generating document:
            
            {str(e)}
            
            Details:
            {error_details}
            """,
            status=500
        )


    finally:
        # Delete temporary generated file
        try:
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
        except Exception:
            pass

# ---------------------------------------------------------------------------
# Bulk export views
# ---------------------------------------------------------------------------

@user_passes_test(lambda u: u.is_staff)

def bulk_export_form(request):
    """Display form for bulk application form export with filters."""
    from accounts.models import StaffRole, College
    from applications.course_mappings import COURSE_COLLEGE_MAPPING
    
    departments = (
        Application.objects.values_list("department", flat=True)
        .distinct()
        .order_by("department")
    )
    years_of_study = (
        Application.objects.values_list("year_of_study", flat=True)
        .distinct()
        .order_by("year_of_study")
    )
    
    # Get college choices
    colleges = College.choices
    
    # Check if user is coordinator
    is_coordinator = request.user.staff_role == StaffRole.COORDINATOR
    assigned_college = request.user.assigned_college if is_coordinator else None

    # Get filtered applications for preview
    filtered_applications = get_filtered_applications(request)

    context = {
        "departments": departments,
        "years_of_study": years_of_study,
        "status_choices": ApplicationStatus.choices,
        "colleges": colleges,
        "is_coordinator": is_coordinator,
        "assigned_college": assigned_college,
        "filtered_applications": filtered_applications,
    }
    return render(request, "exports/bulk_export_form.html", context)


@user_passes_test(lambda u: u.is_staff)

def bulk_export_zip(request):
    """Export application forms as a ZIP file with individual DOCX files."""
    applications = get_filtered_applications(request)

    if not applications:
        return HttpResponse("No applications match the selected filters.", status=400)

    # Only check payment confirmation for applications that are in the filtered results
    unconfirmed = [
        app
        for app in applications
        if not hasattr(app, "payment") or not app.payment.is_confirmed
    ]
    
    # If there are unconfirmed payments in the filtered results, show error
    if unconfirmed:
        return HttpResponse(
            f"Cannot export application forms. {len(unconfirmed)} application(s) "
            "do not have confirmed payments. Use the payment status filter to export only confirmed applications.",
            status=400,
        )

    # Path to DOCX template
    template_path = os.path.join(
        settings.BASE_DIR,
        "templates",
        "exports",
        "application_template.docx"
    )

    # Check template exists
    if not os.path.exists(template_path):
        return HttpResponse(
            "Template file not found. Please create application_template.docx "
            "inside templates/exports/",
            status=500
        )

    # Create temporary directory
    export_dir = os.path.join(settings.BASE_DIR, "tmp_exports")
    os.makedirs(export_dir, exist_ok=True)

    zip_buffer = io.BytesIO()
    temp_files = []

    try:
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for app in applications:
                file_id = uuid.uuid4()
                temp_docx = os.path.join(
                    export_dir,
                    f"application_{file_id}.docx"
                )
                temp_files.append(temp_docx)

                # Load template
                template = DocxTemplate(template_path)

                # Prepare data
                context = {
                    "first_name": app.first_name,
                    "middle_name": app.middle_name or "",
                    "last_name": app.last_name,
                    "full_name": app.full_name,
                    "nationality": app.nationality,
                    "date_of_birth": (
                        app.date_of_birth.strftime("%d/%m/%Y")
                        if app.date_of_birth
                        else ""
                    ),
                    "age": app.age,
                    "email": app.email,
                    "phone_number": app.phone_number,
                    "address": app.address,
                    "city": app.city,
                    "university": app.university,
                    "department": app.department,
                    "course": app.course,
                    "year_of_study": (
                        f"{app.year_of_study} year"
                        if app.year_of_study
                        else ""
                    ),
                    "date_of_admission": app.admission_date_formatted,
                    "year_of_graduation": app.year_of_graduation,
                    "application_date": (
                        app.created_at.strftime("%d/%m/%Y")
                    ),
                    "status": app.get_status_display(),
                    "reference_number": (
                        app.payment.reference_number
                        if hasattr(app, "payment")
                        else "N/A"
                    ),
                }

                # Fill template placeholders
                template.render(context)

                # Save completed document
                template.save(temp_docx)

                # Add to ZIP
                filename = f"{app.full_name.replace(' ', '_')}_Application.docx"
                with open(temp_docx, "rb") as f:
                    zip_file.writestr(filename, f.read())

        zip_buffer.seek(0)

        response = HttpResponse(zip_buffer.getvalue(), content_type="application/zip")
        response["Content-Disposition"] = (
            f'attachment; filename="IET_Application_Forms_'
            f'{datetime.now().strftime("%Y%m%d_%H%M%S")}.zip"'
        )
        return response

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        return HttpResponse(
            f"Error generating ZIP export: {str(e)}\n\nDetails:\n{error_details}",
            status=500
        )

    finally:
        # Clean up temporary files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception:
                pass


@user_passes_test(lambda u: u.is_staff)

def bulk_export_pdf(request):
    """Export application forms as a combined PDF with a cover page."""
    applications = get_filtered_applications(request)

    if not applications:
        return HttpResponse("No applications match the selected filters.", status=400)

    unconfirmed = [
        app
        for app in applications
        if not hasattr(app, "payment") or not app.payment.is_confirmed
    ]
    if unconfirmed:
        return HttpResponse(
            f"Cannot export application forms. {len(unconfirmed)} application(s) "
            "do not have confirmed payments.",
            status=400,
        )

    # Path to DOCX template
    template_path = os.path.join(
        settings.BASE_DIR,
        "templates",
        "exports",
        "application_template.docx"
    )

    # Check template exists
    if not os.path.exists(template_path):
        return HttpResponse(
            "Template file not found. Please create application_template.docx "
            "inside templates/exports/",
            status=500
        )

    # Create temporary directory
    export_dir = os.path.join(settings.BASE_DIR, "tmp_exports")
    os.makedirs(export_dir, exist_ok=True)

    temp_files = []
    combined_docx_path = os.path.join(export_dir, f"combined_{uuid.uuid4()}.docx")

    try:
        # Create combined document
        from docx import Document
        combined_doc = Document()

        # Add cover page
        combined_doc.add_heading('IET Membership Application Forms Export', 0)
        combined_doc.add_paragraph(f'Export Format: PDF (Note: PDF export requires LibreOffice on server)')
        combined_doc.add_paragraph(f'Total Application Forms: {len(applications)}')
        combined_doc.add_paragraph(f'Exported by: {request.user.get_full_name() or request.user.email}')
        combined_doc.add_paragraph(f'Date: {datetime.now().strftime("%d %b %Y")}')
        combined_doc.add_heading('Application List', 1)
        for i, app in enumerate(applications):
            combined_doc.add_paragraph(f"{i+1}. {app.full_name}")
        
        combined_doc.add_paragraph('\n\nNote: Full PDF export requires LibreOffice installed on the server. This is a DOCX file that can be manually converted to PDF.')
        
        # Add page break
        combined_doc.add_page_break()

        # Generate DOCX files for each application and merge
        for app in applications:
            file_id = uuid.uuid4()
            temp_docx = os.path.join(
                export_dir,
                f"application_{file_id}.docx"
            )
            temp_files.append(temp_docx)

            # Load template
            template = DocxTemplate(template_path)

            # Prepare data
            context = {
                "first_name": app.first_name,
                "middle_name": app.middle_name or "",
                "last_name": app.last_name,
                "full_name": app.full_name,
                "nationality": app.nationality,
                "date_of_birth": (
                    app.date_of_birth.strftime("%d/%m/%Y")
                    if app.date_of_birth
                    else ""
                ),
                "age": app.age,
                "email": app.email,
                "phone_number": app.phone_number,
                "address": app.address,
                "city": app.city,
                "university": app.university,
                "department": app.department,
                "course": app.course,
                "year_of_study": (
                    f"{app.year_of_study} year"
                    if app.year_of_study
                    else ""
                ),
                "date_of_admission": app.admission_date_formatted,
                "year_of_graduation": app.year_of_graduation,
                "application_date": (
                    app.created_at.strftime("%d/%m/%Y")
                ),
                "status": app.get_status_display(),
                "reference_number": (
                    app.payment.reference_number
                    if hasattr(app, "payment")
                    else "N/A"
                ),
            }

            # Fill template placeholders
            template.render(context)

            # Save completed document
            template.save(temp_docx)

            # Load the generated document and append to combined document
            app_doc = Document(temp_docx)
            
            # Add heading for this application
            combined_doc.add_heading(f'Application: {app.full_name}', 1)
            
            # Copy all elements from the app document to combined document
            for element in app_doc.element.body:
                combined_doc.element.body.append(element)
            
            # Add page break between applications
            combined_doc.add_page_break()

        # Save combined document
        combined_doc.save(combined_docx_path)

        # Read and return the combined document as DOCX (since PDF requires LibreOffice)
        with open(combined_docx_path, 'rb') as f:
            docx_data = f.read()

        response = HttpResponse(docx_data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response["Content-Disposition"] = (
            f'attachment; filename="IET_Application_Forms_Combined_'
            f'{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
        )
        return response

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        return HttpResponse(
            f"Error generating export: {str(e)}\n\nDetails:\n{error_details}",
            status=500
        )

    finally:
        # Clean up temporary files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception:
                pass
        try:
            if os.path.exists(combined_docx_path):
                os.remove(combined_docx_path)
        except Exception:
            pass


@user_passes_test(lambda u: u.is_staff)

def bulk_export_word(request):
    """Export application forms as a combined Word document with a summary page."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    applications = get_filtered_applications(request)

    if not applications:
        return HttpResponse("No applications match the selected filters.", status=400)

    unconfirmed = [
        app
        for app in applications
        if not hasattr(app, "payment") or not app.payment.is_confirmed
    ]
    if unconfirmed:
        return HttpResponse(
            f"Cannot export application forms. {len(unconfirmed)} application(s) "
            "do not have confirmed payments.",
            status=400,
        )

    # Path to DOCX template
    template_path = os.path.join(
        settings.BASE_DIR,
        "templates",
        "exports",
        "application_template.docx"
    )

    # Check template exists
    if not os.path.exists(template_path):
        return HttpResponse(
            "Template file not found. Please create application_template.docx "
            "inside templates/exports/",
            status=500
        )

    # Create temporary directory
    export_dir = os.path.join(settings.BASE_DIR, "tmp_exports")
    os.makedirs(export_dir, exist_ok=True)

    temp_files = []
    combined_docx_path = os.path.join(export_dir, f"combined_{uuid.uuid4()}.docx")

    try:
        # Create combined document
        combined_doc = Document()

        # Create professional summary page
        title = combined_doc.add_heading('IET Membership Application Forms', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add summary table
        summary_para = combined_doc.add_paragraph()
        summary_para.add_run('Export Summary').bold = True
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create table for summary
        table = combined_doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Set header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '#'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Department'
        hdr_cells[3].text = 'Course'
        
        # Add data rows
        for i, app in enumerate(applications):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = app.full_name
            row_cells[2].text = app.department
            row_cells[3].text = app.course
        
        # Add export metadata
        combined_doc.add_paragraph()
        meta_para = combined_doc.add_paragraph()
        meta_para.add_run('Export Details:').bold = True
        combined_doc.add_paragraph(f'Total Applications: {len(applications)}')
        combined_doc.add_paragraph(f'Exported by: {request.user.get_full_name() or request.user.email}')
        combined_doc.add_paragraph(f'Date: {datetime.now().strftime("%d %b %Y %H:%M")}')
        
        # Add page break before first form
        combined_doc.add_page_break()

        # Generate DOCX files for each application and merge using Composer
        try:
            from docxcompose.composer import Composer
            
            for app in applications:
                file_id = uuid.uuid4()
                temp_docx = os.path.join(
                    export_dir,
                    f"application_{file_id}.docx"
                )
                temp_files.append(temp_docx)

                # Load template
                template = DocxTemplate(template_path)

                # Prepare data
                context = {
                    "first_name": app.first_name,
                    "middle_name": app.middle_name or "",
                    "last_name": app.last_name,
                    "full_name": app.full_name,
                    "nationality": app.nationality,
                    "date_of_birth": (
                        app.date_of_birth.strftime("%d/%m/%Y")
                        if app.date_of_birth
                        else ""
                    ),
                    "age": app.age,
                    "email": app.email,
                    "phone_number": app.phone_number,
                    "address": app.address,
                    "city": app.city,
                    "university": app.university,
                    "department": app.department,
                    "course": app.course,
                    "year_of_study": (
                        f"{app.year_of_study} year"
                        if app.year_of_study
                        else ""
                    ),
                    "date_of_admission": app.admission_date_formatted,
                    "year_of_graduation": app.year_of_graduation,
                    "application_date": (
                        app.created_at.strftime("%d/%m/%Y")
                    ),
                    "status": app.get_status_display(),
                    "reference_number": (
                        app.payment.reference_number
                        if hasattr(app, "payment")
                        else "N/A"
                    ),
                }

                # Fill template placeholders
                template.render(context)

                # Save completed document
                template.save(temp_docx)
            
            # Use Composer to merge documents while preserving formatting
            composer = Composer(combined_doc)
            for temp_docx in temp_files:
                app_doc = Document(temp_docx)
                composer.append(app_doc)
            
            composer.save(combined_docx_path)
            
        except ImportError:
            # Fallback if docxcompose is not available
            for app in applications:
                file_id = uuid.uuid4()
                temp_docx = os.path.join(
                    export_dir,
                    f"application_{file_id}.docx"
                )
                temp_files.append(temp_docx)

                # Load template
                template = DocxTemplate(template_path)

                # Prepare data
                context = {
                    "first_name": app.first_name,
                    "middle_name": app.middle_name or "",
                    "last_name": app.last_name,
                    "full_name": app.full_name,
                    "nationality": app.nationality,
                    "date_of_birth": (
                        app.date_of_birth.strftime("%d/%m/%Y")
                        if app.date_of_birth
                        else ""
                    ),
                    "age": app.age,
                    "email": app.email,
                    "phone_number": app.phone_number,
                    "address": app.address,
                    "city": app.city,
                    "university": app.university,
                    "department": app.department,
                    "course": app.course,
                    "year_of_study": (
                        f"{app.year_of_study} year"
                        if app.year_of_study
                        else ""
                    ),
                    "date_of_admission": app.admission_date_formatted,
                    "year_of_graduation": app.year_of_graduation,
                    "application_date": (
                        app.created_at.strftime("%d/%m/%Y")
                    ),
                    "status": app.get_status_display(),
                    "reference_number": (
                        app.payment.reference_number
                        if hasattr(app, "payment")
                        else "N/A"
                    ),
                }

                # Fill template placeholders
                template.render(context)

                # Save completed document
                template.save(temp_docx)
                
                # Add page break and heading for this application
                combined_doc.add_page_break()
                combined_doc.add_heading(f'Application: {app.full_name}', level=1)
                
                # Load the generated document and copy its body content
                app_doc = Document(temp_docx)
                for element in app_doc.element.body:
                    combined_doc.element.body.append(element)

            combined_doc.save(combined_docx_path)

        # Read and return the combined document
        with open(combined_docx_path, 'rb') as f:
            docx_data = f.read()

        response = HttpResponse(docx_data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response["Content-Disposition"] = (
            f'attachment; filename="IET_Application_Forms_Combined_'
            f'{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
        )
        return response

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        return HttpResponse(
            f"Error generating Word export: {str(e)}\n\nDetails:\n{error_details}",
            status=500
        )

    finally:
        # Clean up temporary files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception:
                pass
        try:
            if os.path.exists(combined_docx_path):
                os.remove(combined_docx_path)
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def get_filtered_applications(request):
    """Get applications based on filter parameters in the request."""
    from applications.course_mappings import COURSE_COLLEGE_MAPPING
    
    queryset = Application.objects.all()

    # Check both GET and POST for filter parameters (POST used for export, GET for filtering)
    request_data = request.POST if request.method == 'POST' else request.GET
    
    department = request_data.get("department")
    year_of_study = request_data.get("year_of_study")
    status = request_data.get("status")
    date_from = request_data.get("date_from")
    date_to = request_data.get("date_to")
    payment_status = request_data.get("payment_status")
    college = request_data.get("college")

    if college:
        queryset = queryset.filter(course__in=[
            course for course, col in COURSE_COLLEGE_MAPPING.items()
            if col == college
        ])
    if department:
        queryset = queryset.filter(department=department)
    if year_of_study:
        queryset = queryset.filter(year_of_study=year_of_study)
    if status:
        queryset = queryset.filter(status=status)
    if date_from:
        queryset = queryset.filter(created_at__gte=date_from)
    if date_to:
        queryset = queryset.filter(created_at__lte=date_to)
    if payment_status:
        if payment_status == "confirmed":
            queryset = queryset.filter(payment__is_confirmed=True)
        elif payment_status == "pending":
            queryset = queryset.filter(
                payment__is_confirmed=False, payment__is_rejected=False
            )
        elif payment_status == "rejected":
            queryset = queryset.filter(payment__is_rejected=True)

    return queryset.select_related("payment", "user")


def generate_cover_page(request, applications, format_type):
    """Generate cover page content for bulk exports."""
    filters = {
        "Department": request.GET.get("department", "All"),
        "Year of Study": request.GET.get("year_of_study", "All"),
        "Status": request.GET.get("status", "All"),
        "Payment Status": request.GET.get("payment_status", "All"),
        "Date From": request.GET.get("date_from", "All"),
        "Date To": request.GET.get("date_to", "All"),
    }

    filter_summary = [
        f"{key}: {value}" for key, value in filters.items() if value and value != "All"
    ]

    name_list = "\n".join(
        [f"{i + 1}. {app.full_name}" for i, app in enumerate(applications)]
    )

    cover_content = f"""
IET Membership Application Forms Export
{'=' * 50}

Export Summary:
{'-' * 30}
Export Format: {format_type}
Total Application Forms: {len(applications)}
Exported by: {request.user.get_full_name() or request.user.email}
Date: {datetime.now().strftime('%d %b %Y')}

Applied Filters:
{'-' * 30}
{chr(10).join(filter_summary) if filter_summary else 'No filters applied'}

Application List:
{'-' * 30}
{name_list}
"""

    return cover_content